from transformers import AutoModelForCausalLM, AutoTokenizer, Trainer, TrainingArguments, DataCollatorForLanguageModeling
from datasets import Dataset
import torch
from docx import Document

# === 1. Create dummy email dataset ===
data = {
    "prompt": [
        "Write a formal email declining a meeting.",
        "Write a follow-up email after a networking event.",
    ],
    "response": [
        "Dear [Name], Thank you for the invitation, but I won’t be able to attend...",
        "Hi [Name], It was a pleasure meeting you at [event]. I’d love to stay in touch...",
    ]
}

def parse_prompts_responses(text_list):
    prompts = []
    responses = []
    current_prompt = ""
    current_response = ""

    for line in text_list:
        if line.lower().startswith("prompt:"):
            current_prompt = line[len("prompt:"):].strip()
        elif line.lower().startswith("response:"):
            current_response = line[len("response:"):].strip()
            # Save pair when both are ready
            if current_prompt and current_response:
                prompts.append(current_prompt)
                responses.append(current_response)
                current_prompt = ""
                current_response = ""

    return {"prompt": prompts, "response": responses}

def load_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():  # Ignore empty lines
            full_text.append(para.text.strip())
    return full_text

from docx import Document
import os

file_path = "D:/AI/myai-learn/emaildataset.docx"  # Use forward slashes or raw string

if os.path.exists(file_path):
    doc = Document(file_path)
    print("Loaded successfully!")
    for para in doc.paragraphs:
        print(para.text)
else:
    print(f"❌ File not found at: {file_path}")

text_list = load_docx(file_path)
parsed_data = parse_prompts_responses(text_list)
dataset = Dataset.from_dict(parsed_data)

# === 2. Choose a lightweight model for CPU ===
model_id = "distilgpt2"  # Much smaller than Mistral

tokenizer = AutoTokenizer.from_pretrained(model_id)
tokenizer.pad_token = tokenizer.eos_token  # Prevent tokenizer padding errors

model = AutoModelForCausalLM.from_pretrained(model_id)



# === 3. Format the dataset for fine-tuning ===
def format_email(example):
    input_text = f"### Instruction:\n{example['prompt']}\n\n### Response:\n{example['response']}"
    return tokenizer(
        input_text,
        truncation=True,
        padding="max_length",
        max_length=256
    )

tokenized_dataset = dataset.map(format_email)

# === 4. Set training arguments ===
training_args = TrainingArguments(
    output_dir="./email-gpt2",
    per_device_train_batch_size=1,
    num_train_epochs=5,
    logging_steps=1,
    save_steps=5,
    save_total_limit=1,
    report_to="none",
    no_cuda=True,  # <== This forces CPU usage
    fp16=False      # <== Don’t use mixed precision on CPU
)

# === 5. Trainer setup ===
data_collator = DataCollatorForLanguageModeling(tokenizer=tokenizer, mlm=False)

trainer = Trainer(
    model=model,
    args=training_args,
    train_dataset=tokenized_dataset,
    data_collator=data_collator,
)

# === 6. Train! ===
trainer.train()

# === 7. Generate a sample email ===
#prompt = "Write a friendly email to a colleague ashutosh about a meeting reschedule."
#input_ids = tokenizer(f"### Instruction:\n{prompt}\n\n### Response:\n", return_tensors="pt").input_ids

prompt="""
Instruction:
You are a helpful AI assistant.

Write a professional out-of-office auto-reply email.

Details:
- I am on leave until 15th April 2025.
- For urgent matters, contact team@example.com.
- The tone should be polite, professional, and to the point.

Response:
"""
inputs = tokenizer(prompt, return_tensors="pt", padding=True)

outputs = model.generate(
    input_ids=inputs["input_ids"],
    attention_mask=inputs["attention_mask"],
    max_new_tokens=150,
    do_sample=False,  # makes output deterministic
    temperature=0.7,
    top_p=0.9,
    repetition_penalty=1.3,
    eos_token_id=tokenizer.eos_token_id
)

response = tokenizer.decode(outputs[0], skip_special_tokens=True)
print(response)